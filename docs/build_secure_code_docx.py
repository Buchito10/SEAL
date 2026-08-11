from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "docs"))

from build_secure_code_appendix_pdf import CONTROLS, read_lines  # noqa: E402


OUT_REPO = ROOT / "output/pdf/Seal - Codificacion Segura con Codigo.docx"
OUT_EXTERNAL = Path(
    "/Users/hugoromero/Documents/Universidad/9no Cuatrimestre/Desarrollo Web Integral/2do Parcial/Seal - Codificacion Segura con Codigo.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "E8F5E9"
YELLOW = "FFF8E1"
RED = "FDECEC"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def cell_margins(table, top=90, start=120, bottom=90, end=120):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def status_fill(status):
    if status == "Cumple":
        return GREEN
    if status == "No cumple":
        return RED
    return YELLOW


def add_inline(paragraph, text, bold=False, color=None):
    parts = text.split("`")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if idx % 2:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(75, 75, 75)
        else:
            run.bold = bold
            if color:
                run.font.color.rgb = color


def para(doc, text="", style=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    add_inline(p, text, bold=bold, color=color)
    return p


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_after = Pt(8)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("SEAL - Codificacion Segura con Codigo")
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED


def add_cover(doc):
    title = para(doc, "SEAL - Codificacion Segura", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = para(doc, "Documento editable para Pages con evidencia de codigo", bold=True, color=BLUE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for line in [
        "Hugo Mauricio Romero Rodriguez",
        "Miguel Angel Leal Perez",
        "Miguel Angel Durazno Martinez",
        "Omar Hernandez Cervantes",
        "DS01SV-25",
        "Hector Saldana Benitez",
        "Desarrollo Web Integral",
        "1 de Julio de 2026",
        "San Juan del Rio, Qro.",
    ]:
        p = para(doc, line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_summary(doc):
    doc.add_heading("Resumen de Cumplimiento", level=1)
    para(
        doc,
        "Este documento resume los controles de codificacion segura evaluados en SEAL y agrega el codigo real que demuestra cada punto. Cuando aplica, se incluye tambien como demostrarlo en la app, terminal o navegador.",
    )
    counts = {"Cumple": 0, "Cumple parcialmente": 0, "No cumple": 0}
    for c in CONTROLS:
        counts[c["status"]] += 1

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    cell_margins(table)
    headers = ["Resultado", "Cantidad", "Puntaje", "Lectura"]
    widths = [2200, 1700, 1700, 3760]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_width(cell, widths[i])
        shade(cell, LIGHT_BLUE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cell.paragraphs[0], h, bold=True, color=DARK_BLUE)
    rows = [
        ("Cumple", str(counts["Cumple"]), "10.0", "Implementado"),
        ("Cumple parcialmente", str(counts["Cumple parcialmente"]), "2.0", "Existe base, falta endurecer"),
        ("No cumple", str(counts["No cumple"]), "0.0", "Pendiente"),
        ("Total", "15", "12/15", "80% aproximado"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_width(cells[i], widths[i])
            if i == 0:
                shade(cells[i], status_fill(value) if value in counts else LIGHT_GRAY)
            add_inline(cells[i].paragraphs[0], value, bold=(row[0] == "Total"))
    doc.add_paragraph()


def add_control_table(doc):
    doc.add_heading("Matriz de Controles", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    cell_margins(table, top=75, bottom=75)
    set_repeat_header(table.rows[0])
    headers = ["#", "Control", "Estado", "Demostracion recomendada"]
    widths = [650, 2900, 1600, 4210]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_width(cell, widths[i])
        shade(cell, LIGHT_BLUE)
        add_inline(cell.paragraphs[0], h, bold=True, color=DARK_BLUE)
    for c in CONTROLS:
        row = table.add_row().cells
        values = [c["id"], c["title"], c["status"], c["demo"]]
        for i, value in enumerate(values):
            cell_width(row[i], widths[i])
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if i == 2:
                shade(row[i], status_fill(c["status"]))
            add_inline(row[i].paragraphs[0], value, bold=(i == 2))
    doc.add_page_break()


def add_code_block(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(7.2)
        run.font.color.rgb = RGBColor(30, 41, 59)


def add_control_details(doc):
    doc.add_heading("Evidencia de Codigo por Punto", level=1)
    for c in CONTROLS:
        doc.add_heading(f"{c['id']} - {c['title']}", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        borders(table)
        cell_margins(table)
        left, right = table.rows[0].cells
        cell_width(left, 1700)
        cell_width(right, 7660)
        shade(left, status_fill(c["status"]))
        shade(right, LIGHT_GRAY)
        add_inline(left.paragraphs[0], f"Estado: {c['status']}", bold=True)
        add_inline(right.paragraphs[0], f"Como demostrarlo: {c['demo']}")
        doc.add_paragraph()
        for path, start, end in c["snippets"]:
            para(doc, f"Archivo: {path} (lineas {start}-{end})", bold=True, color=DARK_BLUE)
            add_code_block(doc, read_lines(path, start, end))
            doc.add_paragraph()


def add_demo_steps(doc):
    doc.add_heading("Guia Rapida para Mostrarlo en la App", level=1)
    steps = [
        "Prender backend: cd Back && npm run dev.",
        "Prender frontend: cd Front && npm run dev.",
        "Entrar a http://localhost:3000/login con cuenta admin.",
        "Mostrar validacion de entradas intentando un login incorrecto o un perfil con datos invalidos.",
        "Mostrar carga de archivos en Admin > Plantillas intentando un archivo no .docx.",
        "Mostrar RBAC entrando con admin y con cliente.",
        "Mostrar JWT/localStorage en DevTools > Application > Local Storage.",
        "Mostrar Helmet con curl -I http://localhost:3001/health.",
        "Mostrar error seguro con curl -i http://localhost:3001/admin/users sin token.",
        "Mostrar logs navegando por la app mientras la terminal del backend esta visible.",
    ]
    for step in steps:
        p = para(doc, step)
        p.style = "List Number"


def build():
    doc = Document()
    configure(doc)
    add_cover(doc)
    add_summary(doc)
    add_control_table(doc)
    add_control_details(doc)
    add_demo_steps(doc)
    add_footer(doc)
    OUT_REPO.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_REPO)
    doc.save(OUT_EXTERNAL)
    print(OUT_REPO)
    print(OUT_EXTERNAL)


if __name__ == "__main__":
    build()
