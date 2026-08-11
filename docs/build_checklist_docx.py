from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "checklist-codificacion-segura-seal.md"
OUTPUT = ROOT / "docs" / "Checklist_Codificacion_Segura_SEAL.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run_with_inline_code(paragraph, text, bold=False, italic=False, color=None):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if color:
                run.font.color.rgb = color
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    add_run_with_inline_code(p, text, bold=bold, italic=italic, color=color)
    return p


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) < 2:
        return

    header = parsed[0]
    body = parsed[2:] if len(parsed) > 2 and all(set(c.replace(":", "").replace("-", "")) == set() for c in parsed[1]) else parsed[1:]
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    set_repeat_table_header(table.rows[0])

    widths = [9360 // cols for _ in range(cols)]
    if cols == 3:
        widths = [1800, 3600, 3960]
    elif cols == 2:
        widths = [2600, 6760]

    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_inline_code(p, text, bold=True, color=DARK_BLUE)

    for row in body:
        cells = table.add_row().cells
        for idx in range(cols):
            value = row[idx] if idx < len(row) else ""
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            if idx == 0 and cols >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_inline_code(p, value)

    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.color.rgb = DARK_BLUE
    title.font.bold = True
    title.paragraph_format.space_after = Pt(4)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("SEAL - Checklist de Codificacion Segura")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def write_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(60, 60, 60)


def add_status_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="C9D7E8")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    add_run_with_inline_code(p, text, bold=True, color=DARK_BLUE)
    doc.add_paragraph()


def build():
    doc = Document()
    configure_styles(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_lines = []

    title_done = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                write_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            i += 1
            if i >= len(lines) or not (lines[i].startswith("|") and lines[i].rstrip().endswith("|")):
                add_markdown_table(doc, table_lines)
                table_lines = []
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not title_done:
                p = add_paragraph(doc, text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_done = True
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("- ["):
            p = add_paragraph(doc, stripped, style=None)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(8)
            continue_i = i + 1
            i = continue_i
            continue

        if stripped.startswith("- "):
            p = add_paragraph(doc, stripped[2:])
            p.style = "List Bullet"
            i += 1
            continue

        if re.match(r"^\\d+\\.\\s+", stripped):
            p = add_paragraph(doc, re.sub(r"^\\d+\\.\\s+", "", stripped))
            p.style = "List Number"
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and ":" in stripped:
            add_status_callout(doc, stripped[2:-2])
            i += 1
            continue

        p = add_paragraph(doc, stripped)
        if stripped.startswith("Total de controles") or stripped.startswith("Puntaje") or stripped.startswith("Cumplimiento"):
            p.runs[0].bold = True
        i += 1

    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
