from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guia_Demostracion_Checklist_SEAL.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"
GREEN_FILL = "E8F5E9"
YELLOW_FILL = "FFF8E1"
RED_FILL = "FDECEC"


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_margins(table, top=80, start=100, bottom=80, end=100):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_text(paragraph, text, bold=False, color=None, size=None):
    parts = text.split("`")
    for idx, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.font.name = "Consolas"
            run.font.size = Pt(8.5 if size is None else size)
            run.font.color.rgb = RGBColor(75, 75, 75)
        else:
            run.bold = bold
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = Pt(size)


def p(doc, text="", style=None, bold=False, color=None):
    para = doc.add_paragraph(style=style)
    add_text(para, text, bold=bold, color=color)
    return para


def bullet(doc, text):
    para = p(doc, text)
    para.style = "List Bullet"
    return para


def numbered(doc, text):
    para = p(doc, text)
    para.style = "List Number"
    return para


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 7, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(21)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE


def status_fill(status):
    if status == "Cumple":
        return GREEN_FILL
    if status == "Parcial":
        return YELLOW_FILL
    return RED_FILL


CONTROLS = [
    {
        "n": "1",
        "block": "Gestion de secretos",
        "control": "Aislamiento de credenciales (.env)",
        "status": "Cumple",
        "does": "Los secretos se leen desde variables de entorno y archivos `.env`; no estan escritos en el codigo.",
        "code": "`Back/src/config/env.js`, `Back/.env.example`, `Front/.env.local.example`.",
        "demo": "Mostrar en codigo `env.js` y explicar que `JWT_SECRET`, `TOKEN_PEPPER`, Firebase, SMTP y Gemini se cargan desde entorno. No conviene abrir valores reales del `.env` en exposicion.",
    },
    {
        "n": "2",
        "block": "Gestion de secretos",
        "control": "Proteccion del repositorio (.gitignore)",
        "status": "Cumple",
        "does": "El repo ignora `.env`, `.env.*`, `node_modules`, `.next`, `Back/secrets`, storage local, llaves y logs.",
        "code": "`.gitignore` en la raiz.",
        "demo": "En terminal ejecutar `git status --ignored` y mostrar que `Back/.env`, `Front/.env.local`, `Front/.next`, `Back/secrets` y `Back/.local-storage` aparecen como ignorados.",
    },
    {
        "n": "3",
        "block": "Gestion de secretos",
        "control": "Entornos separados",
        "status": "Parcial",
        "does": "El sistema usa `NODE_ENV`, URLs de front, storage y variables de entorno, pero no tiene plantillas separadas `.env.development` y `.env.production`.",
        "code": "`Back/src/config/env.js`, `Back/.env.example`, `Front/.env.local.example`.",
        "demo": "Mostrar los `.env.example` y explicar que el siguiente paso es separar archivos por ambiente y bloquear `change-me` en produccion.",
    },
    {
        "n": "4",
        "block": "Validacion y entradas",
        "control": "Validacion estricta de tipos",
        "status": "Cumple",
        "does": "Las peticiones del backend pasan por esquemas Zod antes de la logica de negocio.",
        "code": "`Back/src/validators/*.js`; controladores con `safeParse`.",
        "demo": "En app, intentar login con correo invalido o enviar perfil con CURP/RFC invalido. En codigo, mostrar `auth.schemas.js`, `assignments.schemas.js` o `clientProfile.controller.js`.",
    },
    {
        "n": "5",
        "block": "Validacion y entradas",
        "control": "Saneamiento de entradas",
        "status": "Parcial",
        "does": "Los valores insertados en placeholders se escapan con `escapeHtml`, pero las plantillas HTML completas todavia se renderizan con `dangerouslySetInnerHTML` sin sanitizador formal.",
        "code": "`Back/src/utils/html.js`, `Back/src/utils/placeholderRender.js`; `Front/src/app/(dashboard)/contratos/page.tsx`, `Front/src/app/cliente/dashboard/page.tsx`.",
        "demo": "Mostrar el codigo de `escapeHtml` como evidencia positiva y luego mostrar `dangerouslySetInnerHTML` como pendiente. No hace falta demostrarlo en app; es mejor explicarlo desde codigo.",
    },
    {
        "n": "6",
        "block": "Validacion y entradas",
        "control": "Consultas parametrizadas / ORM",
        "status": "Cumple",
        "does": "No hay SQL concatenado; se usa Firebase Admin SDK con `collection`, `doc`, `where`, `runTransaction`.",
        "code": "`Back/src/services/users.service.js`, `assignments.service.js`, `contracts.service.js`.",
        "demo": "Mostrar servicios Firestore en codigo. Explicar que al no construir SQL con strings se evita SQL Injection clasico.",
    },
    {
        "n": "7",
        "block": "Validacion y entradas",
        "control": "Validacion en carga de archivos",
        "status": "Cumple",
        "does": "La subida de contratos limita tamano, valida `.docx`, MIME, encabezado ZIP, estructura interna y rechaza macros.",
        "code": "`Back/src/controllers/adminContracts.controller.js`, `Back/src/services/docxValidation.service.js`.",
        "demo": "En app Admin > Plantillas intentar subir un archivo no `.docx` o `.docm`. Tambien se puede mostrar el codigo de `validateDocxOrThrow`.",
    },
    {
        "n": "8",
        "block": "Autenticacion y sesiones",
        "control": "Cifrado/hash de contrasenas",
        "status": "Cumple",
        "does": "Las contrasenas se guardan con bcrypt y salt de costo 12; no se guardan en texto plano.",
        "code": "`Back/src/utils/password.js`; usado por login, bootstrap, usuarios y reset.",
        "demo": "Mostrar `bcrypt.genSalt(12)` y `bcrypt.hash`. Si se tiene acceso a Firestore, mostrar que `password_hash` es un hash, no la contrasena.",
    },
    {
        "n": "9",
        "block": "Autenticacion y sesiones",
        "control": "RBAC / menor privilegio",
        "status": "Cumple",
        "does": "Las rutas de admin y cliente exigen JWT y rol. El cliente solo puede ver sus propias asignaciones.",
        "code": "`Back/src/middlewares/authJWT.js`, `requireRole.js`, rutas `admin*.routes.js` y `clientAssignments.routes.js`.",
        "demo": "En app iniciar sesion como admin y luego como cliente: admin entra al dashboard admin, cliente entra a `/cliente/dashboard`. En codigo mostrar `router.use(authJWT, requireRole(\"ADMIN\"))`.",
    },
    {
        "n": "10",
        "block": "Autenticacion y sesiones",
        "control": "Seguridad de JWT",
        "status": "Cumple",
        "does": "El token se firma con `JWT_SECRET`, expira en `JWT_EXPIRES_IN` y solo contiene `userId` y `role`.",
        "code": "`Back/src/utils/token.js`, `Back/src/controllers/auth.controller.js`.",
        "demo": "En app iniciar sesion y revisar DevTools > Application > Local Storage para ver `seal_token`. En codigo mostrar que el payload no contiene password, RFC ni CURP.",
    },
    {
        "n": "11",
        "block": "Autenticacion y sesiones",
        "control": "Cookies seguras",
        "status": "No cumple",
        "does": "La sesion se guarda en `localStorage`; no hay cookies `HttpOnly`, `Secure`, `SameSite`.",
        "code": "`Front/src/lib/auth.ts`.",
        "demo": "En DevTools > Application mostrar `seal_token` y `seal_user` en Local Storage. Explicar que para produccion se recomienda migrar a cookies seguras.",
    },
    {
        "n": "12",
        "block": "Red y encabezados",
        "control": "CORS restrictivo",
        "status": "Parcial",
        "does": "El backend usa CORS, pero `app.use(cors())` no restringe origenes por dominio.",
        "code": "`Back/src/app.js`.",
        "demo": "Mostrar `app.use(cors())` en codigo. Explicar que en produccion deberia usarse `origin: process.env.FRONTEND_ORIGIN`.",
    },
    {
        "n": "13",
        "block": "Red y encabezados",
        "control": "Helmet / encabezados HTTP",
        "status": "Cumple",
        "does": "Express usa `helmet()` para agregar cabeceras de seguridad.",
        "code": "`Back/src/app.js`.",
        "demo": "Con backend prendido ejecutar `curl -I http://localhost:3001/health` y mostrar headers como `X-Content-Type-Options` o CSP generada por Helmet.",
    },
    {
        "n": "14",
        "block": "Errores y logs",
        "control": "Abstraccion de errores",
        "status": "Cumple",
        "does": "El middleware global responde mensajes genericos y no manda stack trace al cliente.",
        "code": "`Back/src/middlewares/errorHandler.js`; utilidades `ok` y `fail`.",
        "demo": "Hacer una peticion sin token a un endpoint protegido y mostrar `401`. Para errores internos, mostrar el middleware en codigo.",
    },
    {
        "n": "15",
        "block": "Errores y logs",
        "control": "Logs de seguridad centralizados",
        "status": "Parcial",
        "does": "Hay logs con `morgan(\"dev\")` y `console.error`, pero no hay logging centralizado/estructurado para produccion.",
        "code": "`Back/src/app.js`, `Back/src/middlewares/errorHandler.js`, `Back/src/services/gemini.service.js`.",
        "demo": "Con backend prendido, navegar la app y mostrar en terminal los logs HTTP de Morgan. Explicar que falta Winston/Pino/Sentry/Cloud Logging.",
    },
]


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    set_margins(table)
    widths = [1700, 2800, 1700, 3160]
    headers = ["Resultado", "Cantidad", "Puntaje", "Lectura"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_shading(cell, LIGHT_BLUE)
        set_width(cell, widths[i])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cell.paragraphs[0], h, bold=True, color=DARK_BLUE)
    rows = [
        ["Cumple", "10", "10.0", "Implementado"],
        ["Cumple parcialmente", "4", "2.0", "Existe base, falta endurecer"],
        ["No cumple", "1", "0.0", "Pendiente"],
        ["Total", "15", "12/15", "80% aproximado"],
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, value in enumerate(r):
            set_width(cells[i], widths[i])
            if i == 0:
                fill = GREEN_FILL if value == "Cumple" else YELLOW_FILL if value == "Cumple parcialmente" else RED_FILL if value == "No cumple" else LIGHT_GRAY
                set_shading(cells[i], fill)
            add_text(cells[i].paragraphs[0], value, bold=(r[0] == "Total"))
    doc.add_paragraph()


def add_matrix(doc):
    doc.add_heading("Resumen de los 15 Controles", level=1)
    p(doc, "Esta tabla sirve para mostrar rapidamente el estado de cada punto. El detalle de codigo y demostracion esta en la seccion siguiente.")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    set_margins(table, top=70, bottom=70, start=80, end=80)
    repeat_header(table.rows[0])
    widths = [600, 2500, 4450, 1810]
    headers = ["#", "Bloque", "Control", "Estado"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_width(cell, widths[i])
        set_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cell.paragraphs[0], h, bold=True, color=DARK_BLUE, size=9)
    for c in CONTROLS:
        cells = table.add_row().cells
        values = [c["n"], c["block"], c["control"], c["status"]]
        for i, value in enumerate(values):
            set_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(cells[i].paragraphs[0], value, bold=True, size=9)
            elif i == 3:
                set_shading(cells[i], status_fill(value))
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(cells[i].paragraphs[0], value, bold=True, size=9)
            else:
                add_text(cells[i].paragraphs[0], value, size=9)
    doc.add_paragraph()


def add_control_details(doc):
    doc.add_heading("Detalle por Control", level=1)
    for c in CONTROLS:
        doc.add_heading(f"{c['n']}. {c['control']} - {c['status']}", level=2)
        bullet(doc, f"Que hace SEAL: {c['does']}")
        bullet(doc, f"Donde se ve en codigo: {c['code']}")
        bullet(doc, f"Como mostrarlo: {c['demo']}")


def add_run_guide(doc):
    doc.add_heading("Como Prender el Backend y el Frontend", level=1)
    p(doc, "Usa dos terminales: una para el backend y otra para el frontend. Si ya tienes dependencias instaladas, puedes omitir `npm install`.")

    doc.add_heading("Terminal 1 - Backend API", level=2)
    for step in [
        "cd \"/Users/hugoromero/Documents/Universidad/8vo Cuatrimestre/Desarrollo Web Profesional/3er Parcial/Seal/Back\"",
        "npm install",
        "npm run dev",
    ]:
        p(doc, step)
    p(doc, "Debe mostrar algo como: `Server running on port 3001`. Prueba salud del API con `http://localhost:3001/health`.")

    doc.add_heading("Terminal 2 - Frontend Next.js", level=2)
    for step in [
        "cd \"/Users/hugoromero/Documents/Universidad/8vo Cuatrimestre/Desarrollo Web Profesional/3er Parcial/Seal/Front\"",
        "npm install",
        "npm run dev",
    ]:
        p(doc, step)
    p(doc, "Abre la aplicacion en `http://localhost:3000/login`.")

    doc.add_heading("Si no tienes usuario administrador", level=2)
    p(doc, "Con el `.env` del backend configurado, puedes crear el admin inicial con:")
    p(doc, "cd \"/Users/hugoromero/Documents/Universidad/8vo Cuatrimestre/Desarrollo Web Profesional/3er Parcial/Seal/Back\"")
    p(doc, "node src/bootstrapAdmin.js")
    p(doc, "Usa el correo y password definidos en `ADMIN_BOOTSTRAP_EMAIL` y `ADMIN_BOOTSTRAP_PASSWORD`.")


def add_demo_script(doc):
    doc.add_heading("Guion Recomendado para Exponer", level=1)
    numbered(doc, "Abrir el documento y mostrar el resumen: 15 controles, 10 cumplen, 4 parciales y 1 no cumple.")
    numbered(doc, "Levantar backend y frontend. En backend dejar visible la terminal para mostrar logs de Morgan.")
    numbered(doc, "Abrir `http://localhost:3000/login` y entrar como administrador.")
    numbered(doc, "Mostrar en app: Dashboard admin, Clientes, Plantillas y Contratos.")
    numbered(doc, "Demostrar validacion de entradas con un login o perfil invalido.")
    numbered(doc, "Demostrar carga de archivos desde Plantillas, idealmente intentando un archivo invalido y explicando la validacion `.docx`.")
    numbered(doc, "Demostrar RBAC iniciando como admin y despues como cliente, o mostrar rutas protegidas en codigo.")
    numbered(doc, "Demostrar JWT y localStorage en DevTools > Application > Local Storage (`seal_token`, `seal_user`). Aclarar que este punto explica tambien por que cookies seguras esta como No cumple.")
    numbered(doc, "Demostrar Helmet con `curl -I http://localhost:3001/health` y los headers de seguridad.")
    numbered(doc, "Cerrar con los pendientes: sanitizacion HTML completa, CORS restrictivo, cookies seguras, entornos separados y logs centralizados.")

    doc.add_heading("Que no es necesario demostrar en la app", level=2)
    bullet(doc, "`.env`, `.gitignore` y entornos separados: mejor demostrarlos en codigo/terminal, no en la app.")
    bullet(doc, "bcrypt: mejor demostrarlo en codigo; no conviene exponer datos reales de Firestore.")
    bullet(doc, "Firestore SDK/consultas parametrizadas: se explica en codigo, porque no es algo visible para el usuario final.")
    bullet(doc, "Sanitizacion parcial: se explica mejor comparando `escapeHtml` contra los usos de `dangerouslySetInnerHTML`.")


def add_commands(doc):
    doc.add_heading("Comandos Utiles para la Demostracion", level=1)
    commands = [
        ("Ver archivos ignorados", "git status --ignored"),
        ("Probar salud del backend", "curl http://localhost:3001/health"),
        ("Ver headers de Helmet", "curl -I http://localhost:3001/health"),
        ("Probar endpoint protegido sin token", "curl -i http://localhost:3001/admin/users"),
        ("Buscar validaciones Zod", "rg -n \"safeParse|z.object\" Back/src"),
        ("Buscar RBAC", "rg -n \"requireRole|authJWT\" Back/src/routes Back/src/middlewares"),
        ("Buscar localStorage", "rg -n \"localStorage\" Front/src Front/components"),
        ("Buscar CORS/Helmet", "rg -n \"cors\\(|helmet\\(\" Back/src/app.js"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    set_margins(table)
    for i, h in enumerate(["Objetivo", "Comando"]):
        cell = table.rows[0].cells[i]
        set_width(cell, [2600, 6760][i])
        set_shading(cell, LIGHT_BLUE)
        add_text(cell.paragraphs[0], h, bold=True, color=DARK_BLUE)
    for label, cmd in commands:
        cells = table.add_row().cells
        set_width(cells[0], 2600)
        set_width(cells[1], 6760)
        add_text(cells[0].paragraphs[0], label)
        add_text(cells[1].paragraphs[0], f"`{cmd}`")


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SEAL - Guia de demostracion del checklist de seguridad")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY


def build():
    doc = Document()
    configure(doc)

    title = p(doc, "Guia de Demostracion del Checklist de Codificacion Segura - SEAL", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "Proyecto: SEAL - Plataforma de gestion, asignacion y firma digital de contratos laborales")
    p(doc, "Alumnos/Equipo: ____________________________________________")
    p(doc, "Fecha de evaluacion: 01 / 07 / 2026")

    doc.add_heading("Objetivo del Documento", level=1)
    p(doc, "Este documento sirve como entregable y como guion de exposicion. Para cada control del checklist indica que implementa SEAL, donde se puede ver en el codigo y como conviene demostrarlo en la app, terminal o repositorio.")

    doc.add_heading("Resumen Ejecutivo", level=1)
    add_summary_table(doc)

    add_run_guide(doc)
    add_demo_script(doc)
    add_matrix(doc)
    add_control_details(doc)
    add_commands(doc)
    add_footer(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
