from __future__ import annotations

import copy
import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path("/Users/hugoromero/Downloads/Minutasdocx.docx")
OUTPUT_DIR = ROOT / "Entregables_SEAL"
REPORT_PATH = OUTPUT_DIR / "Reporte_de_Proyecto_SEAL.docx"
PLAN_PATH = OUTPUT_DIR / "Plan_de_Pruebas_SEAL_Final.docx"
REFERENCE_SHA256 = "2f4a385b1c2de6e4a5098096b1d930d23c544f5c2462252d7477327c9edd02b6"

BLACK = "1A1A1A"
GOLD_LIGHT = "F4ECD6"
GOLD = "D7BE78"
GRAY = "EDEDED"
GREEN = "DDEBDD"
AMBER = "FFF0CC"
RED = "F4D8D8"


def set_font(run, size=Pt(11), bold=None, italic=None, color=BLACK, name="Times New Roman"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_body_keep_section(doc: Document):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=115, start=115, bottom=115, end=115):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_inches):
    widths = [int(round(width * 1440)) for width in widths_inches]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:color"), "9A9A9A")

    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    table._tbl.replace(old_grid, new_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def ensure_decimal_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("lvlJc", "left")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "300")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, before, after in (
        ("Heading 1", 18, 18, 10),
        ("Heading 2", 14, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc._seal_decimal_num_id = ensure_decimal_numbering(doc)


def add_text(doc, text, *, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_font(first, Pt(size), bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest, Pt(size))
    else:
        run = paragraph.add_run(text)
        set_font(run, Pt(size))
    set_paragraph_keep(paragraph, keep_lines=True)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run, Pt(18 if level == 1 else 14 if level == 2 else 12), bold=True)
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, 1, level=0)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    set_font(run, Pt(10))
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, doc._seal_decimal_num_id)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, Pt(10.5))
    return paragraph


def add_cover(doc: Document, title: str):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_font(p.add_run(title), Pt(18), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_font(p.add_run("SEAL - SISTEMA DE ELABORACIÓN Y AUTOMATIZACIÓN\nLEGAL DE CONTRATOS"), Pt(16), bold=True)

    for line in (
        "UNIVERSIDAD TECNOLÓGICA DE SAN JUAN DEL RÍO",
        "INGENIERÍA EN DESARROLLO Y GESTIÓN DE SOFTWARE",
        "ADMINISTRACIÓN DE PROYECTOS DE TI",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run(line), Pt(14))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

    for label, value in (
        ("PROFESORA:", "Alicia Primero Álvarez"),
        ("GRUPO:", "DS01SV-25"),
        (
            "INTEGRANTES:",
            "Hugo Mauricio Romero Rodríguez\nMiguel Ángel Leal Pérez\nMiguel Ángel Durazno Martínez\nOmar Hernández Cervantes",
        ),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(label), Pt(15), bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(9)
        set_font(p.add_run(value), Pt(13.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("10 de agosto de 2026"), Pt(11))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("San Juan del Río, Qro."), Pt(11))


def add_content_section(doc: Document):
    previous = doc.sections[-1]
    default_header_refs = [
        copy.deepcopy(source_ref)
        for source_ref in previous._sectPr.findall(qn("w:headerReference"))
        if source_ref.get(qn("w:type")) == "default"
    ]
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(1.25)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.different_first_page_header_footer = False
    # LibreOffice no hereda de forma estable el encabezado vinculado en este
    # archivo de origen. Se agrega una referencia explícita al mismo header.
    section.footer.is_linked_to_previous = True
    # Linking the footer can make python-docx rebuild the section properties,
    # so add the explicit header reference only after that operation.
    current_sect_pr = section._sectPr
    for old_ref in current_sect_pr.findall(qn("w:headerReference")):
        current_sect_pr.remove(old_ref)
    for source_ref in default_header_refs:
        current_sect_pr.insert(0, source_ref)
    return section


def ensure_even_page_headers(doc: Document):
    """Create a distinct even-page header part for reliable LO rendering."""
    doc.settings.odd_and_even_pages_header_footer = True
    first = doc.sections[0]
    source = first.header
    target = first.even_page_header
    target.is_linked_to_previous = False

    source_el = source._element
    target_el = target._element
    for child in list(target_el):
        target_el.remove(child)

    rel_map = {}
    for rel in source.part.rels.values():
        rel_map[rel.rId] = target.part.relate_to(
            rel._target, rel.reltype, is_external=rel.is_external
        )
    relationship_attrs = (qn("r:id"), qn("r:embed"), qn("r:link"))
    for child in source_el:
        cloned = copy.deepcopy(child)
        for element in cloned.iter():
            for attr in relationship_attrs:
                old_rid = element.get(attr)
                if old_rid in rel_map:
                    element.set(attr, rel_map[old_rid])
        target_el.append(cloned)

    for section in doc.sections[1:]:
        section.even_page_header.is_linked_to_previous = True


def add_index(doc: Document, items):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("ÍNDICE"), Pt(16), bold=True)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, doc._seal_decimal_num_id)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(7)
        set_font(p.add_run(item), Pt(11.5))
    doc.add_page_break()


def add_table(doc: Document, headers, rows, widths, *, font_size=9.2, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, GOLD_LIGHT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), Pt(font_size), bold=True)

    for row_data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if status_col is not None and idx == status_col:
                val = str(value).upper()
                set_cell_shading(cell, GREEN if "CUMPLE" in val or "APROB" in val or "RESUELTO" in val else AMBER if "ATEN" in val or "PARCIAL" in val else RED)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == status_col else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_font(p.add_run(str(value)), Pt(font_size), bold=(idx == status_col))
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)
    return table


def add_callout(doc: Document, label: str, text: str, fill=GOLD_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(f"{label}: "), Pt(11), bold=True)
    set_font(p.add_run(text), Pt(11))
    set_table_geometry(table, [6.5])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def new_from_reference(output_path: Path):
    if hashlib.sha256(REFERENCE.read_bytes()).hexdigest() != REFERENCE_SHA256:
        raise RuntimeError("El documento de referencia cambió; se requiere una nueva destilación.")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, output_path)
    doc = Document(output_path)
    clear_body_keep_section(doc)
    configure_styles(doc)
    section = doc.sections[0]
    # El arte del encabezado de la plantilla ocupa aproximadamente 1.1 pulgadas.
    # Se reserva espacio real para evitar que títulos o tablas lo invadan.
    section.top_margin = Inches(1.25)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    return doc


def build_report():
    doc = new_from_reference(REPORT_PATH)
    add_cover(doc, "REPORTE DEL PROYECTO")
    add_content_section(doc)
    add_index(
        doc,
        [
            "Resumen ejecutivo",
            "Indicadores clave de rendimiento (KPIs)",
            "Módulos y alcance entregado",
            "Estado del proyecto y calidad",
            "Incidencias, riesgos y acciones",
            "Lecciones aprendidas y trabajo futuro",
            "Conclusión y fuentes de evidencia",
        ],
    )

    add_heading(doc, "1. RESUMEN EJECUTIVO", 1)
    add_text(
        doc,
        "SEAL es una plataforma web para elaborar, administrar, asignar y firmar contratos. La solución integra un frontend Next.js, una API REST con Express, persistencia y autenticación mediante Firebase, generación de documentos y PDF, firma móvil mediante token/QR, mensajería, asistencia con IA y controles de acceso por rol.",
    )
    add_text(
        doc,
        "Al corte del 10 de agosto de 2026, el proyecto dispone de una versión desplegada en una máquina virtual de Azure mediante Docker, Nginx y HTTPS de Let's Encrypt. Las pruebas documentadas cubren lógica crítica de plantillas, integración del inicio de sesión, carga del sitio, accesibilidad y verificación de infraestructura. El resultado general es favorable, con acciones pendientes en cobertura de ramas, rendimiento de frontend, auditoría de dependencias moderadas y ampliación de pruebas end-to-end.",
    )
    add_callout(
        doc,
        "Estado general",
        "EN CONTROL CON ACCIONES DE MEJORA. Los criterios críticos de liberación documentados se aprobaron; no se registraron fallos HTTP en la prueba de carga de producción.",
        GREEN,
    )
    add_text(
        doc,
        "Nota de alcance del reporte: no se recibió una línea base formal de presupuesto ni de cronograma, por lo que no se inventan indicadores de costo o variación de calendario. Los KPIs siguientes se limitan a evidencia verificable del repositorio y de los reportes QA.",
        size=10.5,
    )

    add_heading(doc, "2. INDICADORES CLAVE DE RENDIMIENTO (KPIs)", 1)
    kpis = [
        ("Casos automatizados aprobados", "13/13 (100 %) / meta 100 %", "CUMPLE", "11 Jest + 2 Playwright; RESULTADOS_QA_PRODUCCION.md"),
        ("Cobertura de código seleccionado", "46/46 líneas; 50/51 sent. / meta ≥ 80 %", "CUMPLE", "100 % líneas; 98.03 % sentencias; coverage-summary.json"),
        ("Cobertura de ramas", "19/24 (79.16 %) / meta ≥ 80 %", "ATENCIÓN", "Back/coverage/coverage-summary.json"),
        ("Carga: percentil 95", "290.56 ms / meta < 500 ms", "CUMPLE", "qa/reports/k6/summary-production.json"),
        ("Carga: tasa de error HTTP", "0.00 % / meta < 1 %", "CUMPLE", "qa/reports/k6/summary-production.json"),
        ("Checks de carga", "4,009/4,024 (99.63 %) / meta > 99 %", "CUMPLE", "qa/reports/k6/summary-production.json"),
        ("Accesibilidad Lighthouse", "95/100 / meta ≥ 85", "CUMPLE", "qa/reports/lighthouse/seal-local-http.report.json"),
        ("Mejores prácticas Lighthouse", "96/100 / meta ≥ 85", "CUMPLE", "Mismo reporte Lighthouse"),
        ("Rendimiento Lighthouse", "74/100 / meta ≥ 85", "ATENCIÓN", "FCP 0.9 s; LCP 5.5 s; reporte local"),
        ("Verificación de producción", "HTTPS, /login y /api/health aprobados", "CUMPLE", "docs/release/RESULTADOS_QA_PRODUCCION.md"),
        ("Vulnerabilidades altas/críticas", "0 / meta 0", "CUMPLE", "0 frontend; 0 altas/críticas backend; 8 moderadas pendientes"),
    ]
    add_table(doc, ["Indicador", "Resultado / meta", "Estado", "Evidencia"], kpis, [1.55, 1.55, 0.9, 2.5], font_size=8.5, status_col=2)

    doc.add_page_break()
    add_heading(doc, "3. MÓDULOS Y ALCANCE ENTREGADO", 1)
    modules = [
        ("Autenticación y usuarios", "Inicio de sesión, activación, cambio/restablecimiento de contraseña, invitaciones, perfiles, roles y deshabilitación de usuarios."),
        ("Contratos y plantillas", "Carga de contratos, catálogo de placeholders, versionado, comparación, borradores, publicación, clonación y bloqueo de edición."),
        ("Asignaciones y flujo", "Prevalidación, creación de asignaciones, mensajes, aprobación/rechazo, PDF y seguimiento para administrador y cliente."),
        ("Firma digital", "Generación y validación de tokens, flujo móvil, registro de firma y consulta de evidencia de firma."),
        ("IA y colaboración", "Chats administrativos, edición asistida de plantillas, publicación a contrato y preguntas contextualizadas en asignaciones."),
        ("Interfaz y cumplimiento", "Dashboard, clientes, contratos, plantillas, notificaciones, bitácora, ajustes, política de privacidad y protección de datos."),
        ("Infraestructura y liberación", "Docker Compose, proxy Nginx, HTTPS/Certbot, health checks, scripts de despliegue, verificación, rollback y límites para VM."),
    ]
    add_table(doc, ["Módulo", "Resultado entregado"], modules, [1.65, 4.85], font_size=9.5)
    add_text(doc, "El inventario del repositorio contiene 62 endpoints Express y 16 páginas Next.js. Este conteo describe superficie implementada, no cobertura funcional completa; por ello no se presenta como porcentaje de avance.", size=10.5)

    add_heading(doc, "4. ESTADO DEL PROYECTO Y CALIDAD", 1)
    add_heading(doc, "4.1 Calidad funcional y técnica", 2)
    for item in (
        "Jest aprobó 11 pruebas sobre extracción, validación y renderizado seguro de placeholders, escape de HTML y generación de hash SHA-256.",
        "Playwright aprobó dos escenarios de integración del login: autenticación válida con almacenamiento de sesión/redirección y rechazo 401 sin crear sesión.",
        "El frontend completó lint y compilación de producción durante la revisión del 10 de agosto de 2026.",
        "La cobertura reportada se refiere a tres utilidades críticas seleccionadas; no representa a todo el backend.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "4.2 Rendimiento, disponibilidad y seguridad", 2)
    for item in (
        "k6 ejecutó 50 usuarios virtuales durante 30 segundos en producción: 1,342 solicitudes, 0 % de error HTTP, promedio de 124.91 ms y p95 de 290.56 ms.",
        "Quince verificaciones individuales superaron 500 ms (máximo 515.86 ms), aunque el percentil global y la tasa de checks de 99.63 % cumplieron los umbrales.",
        "La verificación de producción confirmó redirección HTTP a HTTPS, certificado válido, /login HTTP 200 y /api/health con respuesta {ok:true}.",
        "La auditoría reportó cero vulnerabilidades conocidas de producción en frontend y cero altas/críticas en backend; permanecen ocho moderadas transitivas asociadas al ecosistema Firebase.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "5. INCIDENCIAS, RIESGOS Y ACCIONES", 1)
    incidents = [
        ("Nginx conservaba direcciones internas anteriores al recrear contenedores", "Alto: respuestas 502 tras reconstrucción", "Scripts de release y rollback reinician Nginx para resolver nuevamente frontend y backend.", "RESUELTO"),
        ("Lighthouse no abrió Chrome desde macOS contra producción", "Medio: falta medición Lighthouse productiva", "Se conservó el reporte local de los mismos contenedores; repetir desde CI o desde otra estación.", "ATENCIÓN"),
        ("Avisos de seguridad en Next.js", "Alto: exposición potencial de dependencias", "Actualización de Next.js 16.1.6 a 16.3.0 y nueva auditoría.", "RESUELTO"),
        ("Cobertura de ramas de 79.16 %", "Medio: rutas alternativas sin validar", "Agregar casos para ramas restantes y elevar la meta a 85 %.", "ATENCIÓN"),
        ("LCP de 5.5 s y Performance 74", "Medio: experiencia inicial lenta", "Perfilar carga, dividir JavaScript, optimizar fuentes/recursos y repetir Lighthouse.", "ATENCIÓN"),
        ("Ocho dependencias moderadas transitivas", "Medio: deuda técnica de seguridad", "Evaluar actualización mayor de Firebase en rama controlada y ejecutar regresión.", "ATENCIÓN"),
    ]
    add_table(doc, ["Incidencia / riesgo", "Impacto", "Acción", "Estado"], incidents, [1.7, 1.25, 2.65, 0.9], font_size=8.6, status_col=3)

    add_heading(doc, "6. LECCIONES APRENDIDAS Y TRABAJO FUTURO", 1)
    for item in (
        "Conservar evidencia ejecutable en el repositorio permite distinguir resultados comprobados de funcionalidades únicamente implementadas.",
        "Los umbrales globales de carga deben acompañarse de percentiles, máximos y checks individuales para evitar conclusiones incompletas.",
        "El reinicio coordinado del proxy es parte del despliegue cuando los contenedores cambian de dirección interna.",
        "Los datos de prueba y las evidencias deben ser ficticios y no exponer credenciales, contratos reales ni datos personales.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "6.1 Roadmap recomendado", 2)
    doc._seal_decimal_num_id = ensure_decimal_numbering(doc)
    for item in (
        "Ampliar Playwright a gestión de usuarios, contrato, asignación, firma móvil y control de roles.",
        "Elevar la cobertura de ramas a por lo menos 85 % sin presentar la cobertura de utilidades seleccionadas como cobertura total del backend.",
        "Reducir LCP a menos de 2.5 s y obtener Performance Lighthouse ≥ 85 en producción.",
        "Resolver o aceptar formalmente las ocho dependencias moderadas y automatizar auditorías en cada liberación.",
        "Mantener un tablero histórico de KPIs por versión para observar tendencia, no sólo una medición aislada.",
    ):
        add_numbered(doc, item)

    doc.add_page_break()
    add_heading(doc, "7. CONCLUSIÓN Y FUENTES DE EVIDENCIA", 1)
    add_text(
        doc,
        "SEAL cuenta con una base funcional e infraestructura de liberación verificadas. El proyecto cumple los controles críticos documentados de autenticación, lógica segura de plantillas, construcción, despliegue HTTPS y carga. La siguiente iteración debe concentrarse en ampliar la cobertura funcional end-to-end, mejorar el rendimiento visual y cerrar la deuda técnica señalada por cobertura de ramas y dependencias moderadas.",
    )
    for source in (
        "docs/release/RESULTADOS_QA_PRODUCCION.md",
        "docs/release/RESULTADOS_QA_LOCAL.md",
        "Back/coverage/coverage-summary.json y Back/coverage/index.html",
        "qa/reports/playwright/index.html",
        "qa/reports/k6/summary-production.json",
        "qa/reports/lighthouse/seal-local-http.report.json",
        "Front/.next/BUILD_ID (compilación generada durante la revisión)",
    ):
        add_bullet(doc, source)

    doc.core_properties.title = "Reporte del Proyecto SEAL"
    doc.core_properties.subject = "Indicadores, estado, calidad y acciones del proyecto"
    doc.core_properties.author = "Equipo SEAL"
    ensure_even_page_headers(doc)
    doc.save(REPORT_PATH)


def build_test_plan():
    doc = new_from_reference(PLAN_PATH)
    add_cover(doc, "PLAN DE PRUEBAS")
    add_content_section(doc)
    add_index(
        doc,
        [
            "Objetivo, alcance y exclusiones",
            "Estrategia, ambiente y responsables",
            "Criterios de entrada, salida y severidad",
            "Matriz de trazabilidad",
            "Resumen de ejecución",
            "Casos de prueba, resultados y evidencias",
            "Cobertura pendiente y riesgos",
            "Procedimiento de evidencia, incidencias y aprobación",
        ],
    )

    add_heading(doc, "1. OBJETIVO, ALCANCE Y EXCLUSIONES", 1)
    add_heading(doc, "1.1 Objetivo", 2)
    add_text(doc, "Definir y documentar las pruebas que comprueban el comportamiento de SEAL antes de su liberación. Para cumplir la rúbrica, cada caso ejecutado incluye escenario, resultado esperado, resultado obtenido y evidencia localizable.")
    add_heading(doc, "1.2 Alcance probado", 2)
    for item in (
        "Lógica de placeholders y renderizado seguro de contratos.",
        "Integración del formulario de inicio de sesión con respuestas API simuladas.",
        "Disponibilidad básica del frontend y endpoint de salud en producción.",
        "Carga concurrente del login con 50 usuarios virtuales.",
        "Accesibilidad, mejores prácticas, SEO y rendimiento mediante Lighthouse.",
        "Construcción/lint de frontend, HTTPS, contenedores y auditoría de dependencias.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "1.3 Exclusiones y limitaciones", 2)
    add_text(doc, "No se presentan como aprobados los flujos que no cuentan con evidencia reproducible: CRUD completo de usuarios, versionado y bloqueo de contratos, asignación/aprobación, firma móvil real, notificaciones, bitácora, IA con servicio externo y recuperación de contraseña con correo real. Se registran como cobertura pendiente en la sección 7.")
    add_callout(doc, "Regla de honestidad", "La cobertura de 100 % de líneas corresponde únicamente a tres utilidades seleccionadas (46 líneas), no al backend completo.", AMBER)

    add_heading(doc, "2. ESTRATEGIA, AMBIENTE Y RESPONSABLES", 1)
    strategy_rows = [
        ("Unitaria / caja blanca", "Jest", "Funciones de placeholders, escape, render y hash", "Equipo de desarrollo"),
        ("Integración UI", "Playwright", "Login válido e inválido con API interceptada", "QA / frontend"),
        ("Carga", "k6", "GET /login y /api/health con 50 VUs", "QA / infraestructura"),
        ("Calidad web", "Lighthouse", "Performance, accesibilidad, mejores prácticas y SEO", "QA / frontend"),
        ("Infraestructura", "Docker, Nginx, curl", "HTTPS, redirección, salud y contenedores", "Infraestructura"),
        ("Estática / compilación", "ESLint y Next build", "Errores de código y construcción productiva", "Frontend"),
    ]
    add_table(doc, ["Nivel", "Herramienta", "Cobertura", "Responsable"], strategy_rows, [1.25, 1.0, 3.1, 1.15], font_size=8.8)
    add_heading(doc, "2.1 Ambiente y datos", 2)
    add_text(doc, "Local: macOS, Node.js, navegador Chromium/Chrome, Docker Desktop, frontend Next.js, backend Express y Nginx. Producción: Azure VM roco-v4, Ubuntu 22.04, Docker Compose, Nginx y certificado Let's Encrypt en https://4-154-29-114.sslip.io. Los usuarios y credenciales de Playwright son ficticios; no se emplean datos personales reales.")

    add_heading(doc, "3. CRITERIOS DE ENTRADA, SALIDA Y SEVERIDAD", 1)
    criteria_rows = [
        ("Entrada", "Código disponible; dependencias instaladas; ambiente o URL accesible; datos ficticios; caso y meta definidos."),
        ("Salida aprobada", "Resultado obtenido coincide con el esperado, evidencia guardada y umbrales satisfechos."),
        ("Salida fallida", "Diferencia reproducible, incidencia registrada con severidad, evidencia y pasos de repetición."),
        ("Salida bloqueada", "Dependencia externa impide ejecutar; se documenta causa, responsable y nueva fecha."),
        ("Evidencia mínima", "Reporte HTML/JSON, captura, salida de comando o documento firmado por responsable; sin secretos ni datos reales."),
    ]
    add_table(doc, ["Criterio", "Definición"], criteria_rows, [1.45, 5.05], font_size=9.5)
    severity_rows = [
        ("Crítica", "Pérdida/exposición de datos, caída total, firma inválida o acceso no autorizado.", "Bloquea liberación."),
        ("Alta", "Flujo principal inutilizable sin alternativa.", "Corregir y reejecutar antes de liberar."),
        ("Media", "Función degradada con alternativa o impacto limitado.", "Plan de corrección y regresión."),
        ("Baja", "Detalle visual o de texto sin afectar el flujo.", "Puede programarse en backlog."),
    ]
    add_table(doc, ["Severidad", "Definición", "Tratamiento"], severity_rows, [1.0, 3.5, 2.0], font_size=9.2)

    add_heading(doc, "4. MATRIZ DE TRAZABILIDAD", 1)
    trace_rows = [
        ("RNF-SEG-01", "Validación y render seguro de plantillas", "CP-01 a CP-07", "Automatizada"),
        ("RF-AUT-01", "Inicio de sesión y manejo de sesión", "CP-08 y CP-09", "Automatizada"),
        ("RNF-REN-01", "Respuesta bajo concurrencia", "CP-10", "Automatizada"),
        ("RNF-USAB-01", "Accesibilidad y calidad web", "CP-11", "Automatizada"),
        ("RNF-INF-01", "HTTPS, salud y contenedores", "CP-12 y CP-13", "Verificación"),
        ("RNF-CAL-01", "Construcción y calidad estática", "CP-14", "Automatizada"),
        ("RNF-SEG-02", "Riesgo conocido de dependencias", "CP-15", "Auditoría"),
    ]
    add_table(doc, ["Requisito", "Descripción", "Casos", "Tipo"], trace_rows, [1.15, 3.1, 1.2, 1.05], font_size=9.0)

    add_heading(doc, "5. RESUMEN DE EJECUCIÓN", 1)
    summary_rows = [
        ("Jest", "11", "11", "0", "2 suites; 100 % líneas seleccionadas"),
        ("Playwright", "2", "2", "0", "Login válido e inválido"),
        ("k6 producción", "1", "1", "0", "p95 290.56 ms; 0 % error HTTP"),
        ("Lighthouse local", "1", "1 parcial", "0", "Acc. 95 y BP 96; Performance 74 a mejorar"),
        ("Infraestructura", "2", "2", "0", "HTTPS, login, health y contenedores"),
        ("Build/lint", "1", "1", "0", "Next build y ESLint correctos"),
        ("Auditoría dependencias", "1", "1 parcial", "0", "0 altas/críticas; 8 moderadas"),
    ]
    add_table(doc, ["Suite", "Planificados", "Aprobados", "Fallidos", "Observación"], summary_rows, [1.15, 0.85, 0.85, 0.75, 2.9], font_size=8.6)
    add_callout(doc, "Resultado", "15 casos documentados: 13 cumplen completamente y 2 cumplen con observaciones de mejora. No existen casos fallidos en la evidencia conservada.", GREEN)

    doc.add_page_break()
    add_heading(doc, "6. CASOS DE PRUEBA, RESULTADOS Y EVIDENCIAS", 1)
    cases = [
        ("CP-01\nPlaceholders únicos\nEjecutar Jest con tokens repetidos.", "Devuelve employee.name y company.salary sin duplicados.", "APROBADO. La lista contiene ambos tokens una sola vez.", "Back/tests/unit/placeholders.unit.test.js; cobertura HTML"),
        ("CP-02\nCatálogo permitido\nValidar employee.name y company.position.", "Acepta únicamente los placeholders catalogados.", "APROBADO. Devuelve used_placeholders con los dos campos.", "Mismo test unitario; Back/coverage/index.html"),
        ("CP-03\nCampo desconocido\nValidar employee.password.", "Rechaza con error 400 y reporta el campo inválido.", "APROBADO. Se produjo Invalid placeholder(s), código 400.", "Mismo test unitario; reporte Jest"),
        ("CP-04\nEscape HTML\nProbar nulos, script, &, comillas y apóstrofe.", "Convierte caracteres peligrosos y nulos a salida segura.", "APROBADO. Las cinco ramas devolvieron el texto escapado esperado.", "Back/tests/whitebox/template-render.whitebox.test.js"),
        ("CP-05\nRender de plantilla\nSustituir tokens y omitir valor ausente.", "Escapa valores, reemplaza tokens y deja vacío el dato faltante.", "APROBADO. Coincidencia exacta con el HTML esperado.", "Test whitebox; Back/coverage/placeholderRender.js.html"),
        ("CP-06\nRechazo data-ph\nEnviar span con data-ph.", "Rechaza antes del catálogo con DISALLOWED_DATA_PH.", "APROBADO. Error 400 y razón esperada.", "Test whitebox; reporte Jest"),
        ("CP-07\nHash de evidencia\nCalcular SHA-256 de SEAL.", "Hash estable a5fd...fad4.", "APROBADO. Coincidencia exacta del hash.", "Test whitebox; Back/coverage/html.js.html"),
        ("CP-08\nLogin válido\nEnviar qa@seal.test y respuesta API 200 simulada.", "Guarda token/usuario y redirige al Dashboard.", "APROBADO. URL /, Dashboard visible, token y rol ADMIN guardados.", "qa/e2e/login.integration.spec.js; qa/reports/playwright/index.html"),
        ("CP-09\nLogin inválido\nSimular API 401.", "Muestra Credenciales inválidas, conserva /login y no crea token.", "APROBADO. Mensaje visible, URL /login y localStorage vacío.", "Mismo spec y reporte Playwright"),
        ("CP-10\nCarga producción\n50 VUs por 30 s sobre /login.", "Error <1 %, p95 <500 ms y checks >99 %.", "APROBADO. 1,342 req.; 0 % error; p95 290.56 ms; 99.63 % checks.", "qa/reports/k6/summary-production.json"),
        ("CP-11\nLighthouse\nAuditar página de login.", "Accesibilidad y BP ≥85; Performance objetivo ≥85.", "PARCIAL. Acc. 95, BP 96 y SEO 100; Performance 74, LCP 5.5 s.", "qa/reports/lighthouse/seal-local-http.report.json"),
        ("CP-12\nHTTPS y salud\nVerificar redirección, /login y /api/health.", "301 a HTTPS, login 200 y health {ok:true}.", "APROBADO. Los tres resultados se documentaron en producción.", "docs/release/RESULTADOS_QA_PRODUCCION.md"),
        ("CP-13\nContenedores\nInspeccionar frontend, backend y Nginx.", "Frontend/backend saludables y proxy público operativo.", "APROBADO. Servicios saludables; puertos internos 3000/3001 y públicos 80/443.", "RESULTADOS_QA_PRODUCCION.md; compose.yaml"),
        ("CP-14\nBuild y lint\nEjecutar comandos de frontend.", "ESLint sin errores y build genera BUILD_ID.", "APROBADO. Ambos comandos finalizaron correctamente el 10/08/2026.", "Front/.next/BUILD_ID; npm --prefix Front run lint/build"),
        ("CP-15\nDependencias\nRevisar auditoría de producción.", "Cero vulnerabilidades altas o críticas.", "PARCIAL. Meta crítica cumple; permanecen 8 moderadas transitivas de Firebase.", "docs/release/RESULTADOS_QA_LOCAL.md"),
    ]
    add_table(doc, ["ID / caso / procedimiento", "Resultado esperado", "Resultado obtenido", "Evidencia"], cases, [1.75, 1.55, 1.8, 1.4], font_size=7.9, status_col=None)

    add_heading(doc, "7. COBERTURA PENDIENTE Y RIESGOS", 1)
    pending = [
        ("Usuarios", "Crear, invitar, editar, deshabilitar y eliminar; validar roles.", "Alta", "Capturas/API + bitácora"),
        ("Contratos", "Carga DOCX, placeholders, versiones, comparación, borrador, publicación, clon y lock.", "Alta", "Capturas + archivos generados"),
        ("Asignaciones", "Precheck, creación, mensajes, aprobación/rechazo y PDF.", "Alta", "Playwright + PDF"),
        ("Firma móvil", "Token/QR, expiración, firma, evidencia e idempotencia.", "Alta", "Video/capturas + registro"),
        ("Cliente", "Perfil, asignaciones, mensajes, IA y descarga de PDF.", "Media", "Playwright"),
        ("Recuperación", "Solicitud, verificación, expiración y confirmación por correo.", "Alta", "Correo ficticio + API"),
        ("Seguridad", "Rate limit, JWT inválido, acceso cruzado de roles y tamaño de payload.", "Alta", "Pruebas API automatizadas"),
        ("Observabilidad", "Notificaciones, bitácora y registro de incidencias.", "Media", "Capturas + logs"),
    ]
    add_table(doc, ["Área", "Caso pendiente", "Prioridad", "Evidencia requerida"], pending, [1.1, 3.45, 0.85, 1.1], font_size=8.6)
    add_text(doc, "Riesgo principal: los resultados existentes son sólidos para los componentes probados, pero todavía no demuestran regresión integral de todos los flujos de negocio. La liberación académica puede conservarse; una liberación comercial debería cerrar primero los casos de prioridad alta.", size=10.5)

    add_heading(doc, "8. PROCEDIMIENTO DE EVIDENCIA, INCIDENCIAS Y APROBACIÓN", 1)
    add_heading(doc, "8.1 Ejecución y resguardo", 2)
    doc._seal_decimal_num_id = ensure_decimal_numbering(doc)
    for item in (
        "Preparar datos ficticios y registrar versión, ambiente, fecha y responsable.",
        "Ejecutar el comando o procedimiento exacto del caso.",
        "Comparar resultado obtenido contra esperado y asignar Aprobado, Fallido, Bloqueado o Parcial.",
        "Guardar reporte/captura con nombre CP-XX-fecha; ocultar tokens, correos reales y datos contractuales.",
        "Si falla, abrir incidencia con severidad, pasos, evidencia, responsable y versión afectada; después corregir y reejecutar regresión.",
    ):
        add_numbered(doc, item)
    add_heading(doc, "8.2 Comandos reproducibles", 2)
    commands = [
        ("Jest", "npm --prefix Back run test:coverage", "Back/coverage/index.html"),
        ("Playwright", "BASE_URL=https://URL npm --prefix qa run test:e2e", "qa/reports/playwright/index.html"),
        ("k6", "BASE_URL=https://URL k6 run qa/k6/load-test.js", "qa/reports/k6/summary-production.json"),
        ("Lighthouse", "BASE_URL=https://URL npm --prefix qa run lighthouse", "qa/reports/lighthouse/"),
        ("Frontend", "npm --prefix Front run lint && npm --prefix Front run build", "Front/.next/BUILD_ID"),
    ]
    add_table(doc, ["Suite", "Comando", "Salida / evidencia"], commands, [1.0, 3.7, 1.8], font_size=8.7)
    add_heading(doc, "8.3 Aprobación", 2)
    approval = [
        ("Product Owner", "Hugo Mauricio Romero Rodríguez", "____________________"),
        ("Scrum Master", "Miguel Ángel Leal Pérez", "____________________"),
        ("Equipo de desarrollo", "Omar Hernández Cervantes / Miguel Ángel Durazno Martínez", "____________________"),
    ]
    add_table(doc, ["Rol", "Nombre", "Firma"], approval, [1.45, 3.55, 1.5], font_size=9.5)

    doc.core_properties.title = "Plan de Pruebas SEAL"
    doc.core_properties.subject = "Casos de prueba, resultados obtenidos y evidencias"
    doc.core_properties.author = "Equipo SEAL"
    ensure_even_page_headers(doc)
    doc.save(PLAN_PATH)


if __name__ == "__main__":
    build_report()
    build_test_plan()
    print(REPORT_PATH)
    print(PLAN_PATH)
