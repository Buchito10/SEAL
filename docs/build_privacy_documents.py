from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

DOCUMENTS = [
    {
        "source": DOCS / "politica-privacidad-seal.md",
        "output": DOCS / "Politica_Privacidad_SEAL.docx",
        "subtitle": "Aviso integral para usuarios de la plataforma",
        "running": "POLÍTICA DE PRIVACIDAD",
    },
    {
        "source": DOCS / "proteccion-datos-personales-seal.md",
        "output": DOCS / "Proteccion_Datos_Personales_SEAL.docx",
        "subtitle": "Guía operativa y matriz de cumplimiento",
        "running": "LINEAMIENTOS DE PROTECCIÓN DE DATOS",
    },
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(24, 23, 20)
GRAY = RGBColor(89, 89, 89)
GOLD = RGBColor(159, 121, 44)
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF8E8"
GREEN = "E8F5E9"
YELLOW = "FFF8E1"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def paragraph_bottom_rule(paragraph, color="2E74B5", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "10")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, result, end])
    set_run_font(run, size=9, color=GRAY)


def add_header_footer(doc, running, version):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run(f"SEAL  |  {running}")
        set_run_font(run, size=8.5, color=GRAY, bold=True)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(0)
        run = footer.add_run(f"Seal Contratos  ·  Versión {version}  ·  Página ")
        set_run_font(run, size=9, color=GRAY)
        add_field(footer, "PAGE")


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text):
    url_pattern = re.compile(r"https?://\S+")
    cursor = 0
    for match in url_pattern.finditer(text):
        add_inline_markup(paragraph, text[cursor:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    add_inline_markup(paragraph, text[cursor:])


def add_inline_markup(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=GRAY)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    add_inline(paragraph, text)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GOLD)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "9F792C")
    borders.append(left)
    p_pr.append(borders)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "D9E2EC")
        borders.append(node)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_markdown_table(doc, lines):
    parsed = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    header = parsed[0]
    body = parsed[2:]
    widths = [1600, 2860, 3840, 1060] if len(header) == 4 else [9360 // len(header)] * len(header)
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])

    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, size=9, color=DARK_BLUE, bold=True)

    for values in body:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if index == len(values) - 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, GREEN if value == "Cumple" else YELLOW)
            run = paragraph.add_run(value.replace("`", ""))
            # Named override: compact compliance matrix, needed to keep four columns readable.
            set_run_font(run, size=8.5, bold=index == 0 or index == len(values) - 1)

    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def extract_header(lines):
    title = lines[0].removeprefix("# ").strip()
    metadata = []
    index = 1
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("## "):
            break
        match = re.match(r"\*\*(.+?):\*\*\s*(.*)", line.rstrip("  "))
        if match:
            metadata.append((match.group(1), match.group(2)))
        index += 1
    return title, metadata, index


def add_masthead(doc, title, subtitle, metadata):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title.upper())
    set_run_font(run, size=23, color=BLACK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(15)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=14, color=GRAY)

    last = None
    for label, value in metadata:
        last = doc.add_paragraph()
        last.paragraph_format.space_after = Pt(2)
        label_run = last.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, bold=True)
        value_run = last.add_run(value)
        set_run_font(value_run, size=10.5)
    if last is not None:
        last.paragraph_format.space_after = Pt(14)
        paragraph_bottom_rule(last)


def render_markdown(doc, lines, start_index, bullet_id, decimal_id):
    index = start_index
    while index < len(lines):
        raw = lines[index].rstrip()
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("> "):
            add_callout(doc, line[2:])
        elif line.startswith("- "):
            paragraph = doc.add_paragraph()
            set_numbering(paragraph, bullet_id)
            add_inline(paragraph, line[2:])
        elif re.match(r"\d+\.\s", line):
            paragraph = doc.add_paragraph()
            set_numbering(paragraph, decimal_id)
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
        else:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, line)
        index += 1


def build(config):
    lines = config["source"].read_text(encoding="utf-8").splitlines()
    title, metadata, start_index = extract_header(lines)
    version = next((value for label, value in metadata if label == "Versión"), "1.0")

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = config["subtitle"]
    doc.core_properties.author = "Equipo Seal"
    doc.core_properties.keywords = "LFPDPPP, privacidad, datos personales, ARCO, Seal"

    bullet_id = add_numbering(doc, "bullet")
    decimal_id = add_numbering(doc, "decimal")
    add_header_footer(doc, config["running"], version)
    add_masthead(doc, title, config["subtitle"], metadata)
    render_markdown(doc, lines, start_index, bullet_id, decimal_id)
    doc.save(config["output"])
    print(config["output"])


if __name__ == "__main__":
    for document in DOCUMENTS:
        build(document)
